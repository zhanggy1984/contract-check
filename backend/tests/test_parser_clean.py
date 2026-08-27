"""parser 清洗接入单测：extract_pdf/extract_docx 输出已清洗，零宽 PDF 不误判文本层。"""
import unittest
from unittest import mock

from app.parser.docx_parser import extract_docx
from app.parser.pdf_parser import extract_pdf


def _s(*cps: int) -> str:
    return "".join(chr(cp) for cp in cps)


class TestExtractPdfClean(unittest.TestCase):
    def _fake_doc(self, pages_text):
        """构造可迭代的伪 pymupdf doc（每页 get_text 返回固定文本，接受 sort 关键字）。"""
        pages = []
        for t in pages_text:
            p = mock.MagicMock()
            p.get_text.return_value = t
            pages.append(p)
        doc = mock.MagicMock()
        doc.__iter__.side_effect = lambda: iter(pages)   # 每次遍历拿新迭代器（防多次迭代耗尽）
        return doc

    @mock.patch("app.parser.pdf_parser.pymupdf.open")
    def test_zero_width_only_pdf_flagged_scanned(self, m_open):
        # 只有零宽字符的 PDF：页级清洗后为空 → 扫描页（修复：不再误判有文本层）
        m_open.return_value = self._fake_doc([_s(0x200B, 0x200B), _s(0x200C)])
        text, page_texts = extract_pdf("/tmp/x.pdf")
        self.assertEqual(text, "")
        self.assertEqual(page_texts, ["", ""])
        self.assertTrue(any(not t.strip() for t in page_texts), "零宽页应判定为扫描页")

    @mock.patch("app.parser.pdf_parser.pymupdf.open")
    def test_cleans_extracted_text(self, m_open):
        # 全角/零宽/多余空行在提取时按页统一清洗；text 为页级 join（页间单换行）
        m_open.return_value = self._fake_doc(["　甲方Ａ" + _s(0x200B) + "条款\r\n", "\n\n\n乙方"])
        text, page_texts = extract_pdf("/tmp/x.pdf")
        self.assertEqual(page_texts, ["甲方A条款", "乙方"])
        self.assertEqual(text, "甲方A条款\n乙方")
        self.assertFalse(any(not t.strip() for t in page_texts))

    @mock.patch("app.parser.pdf_parser.pymupdf.open")
    def test_mixed_scan_and_text_pages(self, m_open):
        # 混合扫描 PDF：部分页有文本层、部分页为空 → page_texts 按页区分，scanned 只含空页
        m_open.return_value = self._fake_doc(["第一条 标的：服务器。", "", "盖章页"])
        text, page_texts = extract_pdf("/tmp/x.pdf")
        self.assertEqual(page_texts, ["第一条 标的：服务器。", "", "盖章页"])
        self.assertEqual([i for i, t in enumerate(page_texts) if not t.strip()], [1])
        self.assertEqual(text, "第一条 标的：服务器。\n\n盖章页")

    @mock.patch("app.parser.pdf_parser.pymupdf.open")
    def test_get_text_uses_sort(self, m_open):
        # 两栏 PDF 乱序修复：get_text 应带 sort=True（阅读顺序）
        doc = self._fake_doc(["左栏"])
        m_open.return_value = doc
        extract_pdf("/tmp/x.pdf")
        next(iter(doc)).get_text.assert_called_with(sort=True)


class TestExtractDocxClean(unittest.TestCase):
    @mock.patch("app.parser.docx_parser.docx.Document")
    def test_cleans_extracted_text(self, m_doc):
        # 新实现走 iter_inner_content（段落/表格交替块），mock 对象非 Table 走 Paragraph 分支
        m_doc.return_value.iter_inner_content.return_value = [
            mock.MagicMock(text=_s(0xFEFF) + "甲方"),
            mock.MagicMock(text=_s(0x200B) + "乙方"),
            mock.MagicMock(text=""),
        ]
        m_doc.return_value.element.body.iter.return_value = iter(())  # 无文本框
        self.assertEqual(extract_docx("/tmp/x.docx"), "甲方\n乙方")

    @mock.patch("app.parser.docx_parser.docx.Document")
    def test_textbox_content_appended(self, m_doc):
        # 正文文本框（w:txbxContent）内容应被提取追加，页眉页脚刻意不取
        m_doc.return_value.iter_inner_content.return_value = [
            mock.MagicMock(text="甲方"),
        ]
        txbx = mock.MagicMock()
        txbx.iter.return_value = [
            mock.MagicMock(text="备注"),
            mock.MagicMock(text="：附件A"),
        ]
        m_doc.return_value.element.body.iter.return_value = iter([txbx])
        self.assertEqual(extract_docx("/tmp/x.docx"), "甲方\n备注：附件A")


class TestExtractDocxTable(unittest.TestCase):
    def _make_docx(self, build) -> str:
        """用真实 python-docx 生成临时 .docx，返回路径（调用方负责删除）。"""
        import tempfile
        import docx

        d = docx.Document()
        build(d)
        fd, path = tempfile.mkstemp(suffix=".docx")
        import os
        os.close(fd)
        d.save(path)
        return path

    def test_table_content_extracted(self):
        # doc.paragraphs 不含表格内段落（价格表/设备清单关键内容载体），
        # iter_inner_content 按文档顺序取段落+表格块，表格逐 cell 提取
        import os
        import docx

        def build(d: docx.Document):
            d.add_paragraph("甲方：甲公司")
            table = d.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "标的"
            table.cell(0, 1).text = "服务器"
            table.cell(1, 0).text = "金额"
            table.cell(1, 1).text = "10000元"
            d.add_paragraph("第二条 违约责任：违约方赔偿损失。")

        path = self._make_docx(build)
        try:
            text = extract_docx(path)
        finally:
            os.unlink(path)
        for expect in ("甲方：甲公司", "标的", "服务器", "金额", "10000元", "第二条 违约责任：违约方赔偿损失。"):
            self.assertIn(expect, text, f"表格/段落内容应被提取：{expect}")

    def test_textbox_in_real_docx(self):
        # 真实 docx 含文本框：w:txbxContent 内容应被提取（无文本框时 iter 为空不报错）
        import os
        path = self._make_docx(lambda d: d.add_paragraph("甲方：甲公司"))
        try:
            text = extract_docx(path)
        finally:
            os.unlink(path)
        self.assertEqual(text, "甲方：甲公司", "无文本框文档应正常提取段落")


if __name__ == "__main__":
    unittest.main()
